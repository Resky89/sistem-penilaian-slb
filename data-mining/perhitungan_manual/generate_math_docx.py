# -*- coding: utf-8 -*-
"""
Generator Dokumen Panduan Perhitungan Manual Skripsi
Menggunakan 10 data awal + 10 data akhir dari DATA TRAINING.
Semua nilai dihitung dari data real, bukan hardcoded.
"""
import os
import re
import pandas as pd
import numpy as np
import joblib
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.preprocessing import LabelEncoder
from sklearn.ensemble import RandomForestClassifier
from sklearn.multioutput import MultiOutputClassifier
from sklearn.model_selection import GridSearchCV
import shap

# --- PATHS ---
DATA_PATH = r"d:\projects\Sistem-Penilaian-SLB\DATA MENTAH.xlsx"
SPLIT_PATH = r"d:\projects\Sistem-Penilaian-SLB\DATA_SPLIT_TRAIN_TEST.xlsx"
OUTPUT_PATH = r"d:\projects\Sistem-Penilaian-SLB\Panduan_Perhitungan_Manual_Skripsi.docx"

STOPWORDS = [
    'dan', 'di', 'ke', 'dari', 'ini', 'itu', 'juga', 'untuk', 'pada', 'dengan',
    'adalah', 'yang', 'saya', 'kami', 'anda', 'mereka', 'ia', 'dia', 'kita',
    'dapat', 'harus', 'akan', 'sudah', 'telah', 'sedang', 'ingin',
    'ada', 'bukan', 'hanya', 'saja', 'atau', 'namun', 'tetapi',
    'oleh', 'seperti', 'maka', 'jika', 'karena', 'sehingga', 'bahwa',
    'hal', 'secara', 'tersebut', 'dalam', 'atas', 'bawah', 'serta', 'bagi', 'setelah',
    'peserta', 'didik', 'siswa', 'aspek', 'kegiatan', 'aktivitas', 'cara',
    'berikan', 'gunakan', 'saran', 'hari',
    'dafa', 'dzikri', 'eria', 'arif', 'arifin', 'rama', 'fadil', 'afif',
    'azzam', 'dewi', 'faiz', 'ahmad', 'bilqis', 'rizky', 'mia', 'andreansyah',
    'rahman', 'khansa', 'robi', 'roby'
]


# ==================== HELPERS ====================

def clean_text(text):
    if not isinstance(text, str):
        return ""
    text = text.lower()
    text = re.sub(r'[^a-zA-Z\s]', ' ', text)
    words = text.split()
    words = [w for w in words if w not in STOPWORDS and len(w) > 2]
    return " ".join(words)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def add_equation(paragraph, xml_str):
    paragraph._p.append(parse_xml(xml_str))


def fmt(run, size=12, bold=False, italic=False, color=(0, 0, 0), name="Times New Roman"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color)


def build_table(doc, headers, rows_data, col_widths, header_bg="2C3E50",
                font_size=9, center_cols=None, highlight_row_fn=None):
    """Generic table builder to avoid repetition."""
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, header_bg)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt(p.runs[0], size=font_size, bold=True, color=(255, 255, 255))

    # Data rows
    center_cols = center_cols or []
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            set_cell_margins(cell)

            # Separator row highlight
            if highlight_row_fn and highlight_row_fn(r_idx, row_data):
                set_cell_bg(cell, "E8F8F5")
            elif r_idx % 2 == 1:
                set_cell_bg(cell, "F8F9FA")

            # Separator styling
            if str(val).startswith("..."):
                set_cell_bg(cell, "FFF3CD")

            p = cell.paragraphs[0]
            if c_idx in center_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0] if p.runs else p.add_run()
            fmt(run, size=font_size)

    # Column widths
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    return table


# ==================== DATA PREPARATION ====================

def load_training_data():
    """Load training data and prepare features."""
    xl = pd.ExcelFile(SPLIT_PATH)
    train = pd.read_excel(xl, sheet_name=xl.sheet_names[0])

    train['X1 (Nilai)'] = pd.to_numeric(train['X1 (Nilai)'], errors='coerce').fillna(0)
    train['clean_X2'] = train['X2 (Deskripsi Capaian)'].apply(clean_text)

    return train


def get_subset(train):
    """Get first 10 + last 10 rows."""
    head = train.head(10).copy()
    tail = train.tail(10).copy()
    head['_pos'] = 'awal'
    tail['_pos'] = 'akhir'
    subset = pd.concat([head, tail]).reset_index(drop=True)
    # Original index in training set
    orig_indices = list(range(1, 11)) + list(range(len(train) - 9, len(train) + 1))
    subset['_orig_idx'] = orig_indices
    return subset


def compute_tfidf_manual(subset, full_corpus):
    """Compute TF-IDF manually for tracing."""
    N = len(full_corpus)
    results = []

    for i, row in subset.iterrows():
        doc_text = row['clean_X2']
        words = doc_text.split()
        Nd = len(words)
        orig_idx = row['_orig_idx']

        word_freq = {}
        for w in words:
            word_freq[w] = word_freq.get(w, 0) + 1

        # Sort by frequency desc, take top 3
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:3]

        for word, freq in sorted_words:
            df_val = sum(1 for d in full_corpus if word in d.split())
            tf = freq / Nd if Nd > 0 else 0
            idf = np.log(N / df_val) if df_val > 0 else 0
            tfidf_score = tf * idf

            results.append({
                'idx': orig_idx,
                'nama': row['Nama Siswa'],
                'aspek': row['Aspek / Mapel'],
                'word': word,
                'freq': freq,
                'Nd': Nd,
                'tf': round(tf, 4),
                'N': N,
                'df': df_val,
                'idf': round(idf, 4),
                'tfidf': round(tfidf_score, 4)
            })

    return results


# ==================== DOCUMENT GENERATION ====================

def generate_docx():
    print("Loading training data...")
    train = load_training_data()
    subset = get_subset(train)
    N_total = len(train)

    print(f"Training data: {N_total} rows")
    print(f"Subset: {len(subset)} rows (10 awal + 10 akhir)")

    # TF-IDF vectorizer on full training
    tfidf_vec = TfidfVectorizer(max_features=500, ngram_range=(1, 1))
    X_tfidf_full = tfidf_vec.fit_transform(train['clean_X2']).toarray()
    feature_names = list(tfidf_vec.get_feature_names_out())
    n_tfidf_features = len(feature_names)

    # Full feature matrix
    X_nilai_full = train[['X1 (Nilai)']].values
    X_full = np.hstack([X_nilai_full, X_tfidf_full])

    # Encode labels
    le_label = LabelEncoder()
    y_labels = le_label.fit_transform(train['Y (Label)'])

    # Train model for voting & SHAP
    print("Training model for real calculations...")
    rf = RandomForestClassifier(n_estimators=100, max_depth=None, min_samples_split=2, random_state=42)
    rf.fit(X_full, y_labels)

    # Subset features
    head_indices = list(range(10))
    tail_indices = list(range(N_total - 10, N_total))
    subset_indices = head_indices + tail_indices
    X_subset = X_full[subset_indices]

    # Manual TF-IDF calc
    print("Computing manual TF-IDF...")
    tfidf_manual = compute_tfidf_manual(subset, train['clean_X2'].tolist())

    # RF voting
    print("Extracting RF voting...")
    preds_all_trees = np.array([tree.predict(X_subset) for tree in rf.estimators_])

    # SHAP
    print("Computing SHAP values...")
    explainer = shap.TreeExplainer(rf)
    shap_values = explainer.shap_values(X_subset)
    base_values = explainer.expected_value

    # Hyperparameter tuning (real run)
    print("Running GridSearchCV for real HPT log...")
    param_grid = {
        'n_estimators': [50, 100],
        'max_depth': [10, 20, None],
        'min_samples_split': [2, 5]
    }
    grid = GridSearchCV(
        RandomForestClassifier(random_state=42),
        param_grid, cv=3, scoring='accuracy', n_jobs=-1
    )
    grid.fit(X_full, y_labels)

    # Extract top results
    cv_results = pd.DataFrame(grid.cv_results_)
    cv_results = cv_results.sort_values('rank_test_score')
    best_params = grid.best_params_
    best_score = grid.best_score_

    # ==================== BUILD DOCUMENT ====================
    print("Building Word document...")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(4)
        section.left_margin = Cm(4)
        section.bottom_margin = Cm(3)
        section.right_margin = Cm(3)

    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)

    # --- TITLE ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "PANDUAN PENULISAN FORMULA DAN PERHITUNGAN MANUAL\n"
        "UNTUK BAB IV HASIL DAN PEMBAHASAN SKRIPSI"
    )
    fmt(r, size=14, bold=True, color=(44, 62, 80))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Studi Kasus: Penerapan Metode Random Forest Untuk Penilaian Perkembangan\n"
        "Capaian Pembelajaran Siswa SLB Tuna Grahita\n"
        "Muhammad Resky Prabowo Sutejo (NPM: 12522021)"
    )
    fmt(r, size=11, italic=True, color=(127, 140, 141))
    doc.add_paragraph()

    # =====================================================
    # SECTION 1: VARIABEL PENELITIAN
    # =====================================================
    h = doc.add_heading(level=1)
    r = h.add_run("1. Definisi Operasional Variabel Penelitian")
    fmt(r, size=14, bold=True, color=(44, 62, 80))

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(
        "Penelitian ini menerapkan model klasifikasi berbasis machine learning dengan pendekatan "
        "Multi-Output Random Forest. Setiap baris data mewakili penilaian satu aspek/mapel dari seorang siswa. "
        "Model memprediksi label perkembangan (Baik, Cukup, Perlu Bimbingan) berdasarkan "
        "dua variabel independen: Nilai Kuantitatif (X1) dan Deskripsi Capaian (X2). "
        "Multi-output digunakan untuk memprediksi label dari setiap aspek/mapel siswa secara simultan."
    )
    fmt(r)

    # Variabel table
    var_headers = ["No", "Variabel", "Kode & Peran", "Definisi Operasional", "Skala Data"]
    var_data = [
        ["1", "Nilai Kuantitatif", "X1\n(Independen)",
         "Nilai numerik perkembangan capaian belajar siswa pada aspek/mapel terkait (skala 0-100). "
         "Nilai '-' pada data portofolio dikonversi menjadi 0.", "Rasio\n(Kuantitatif)"],
        ["2", "Deskripsi Capaian", "X2\n(Independen)",
         f"Uraian deskriptif kualitatif perkembangan kemampuan siswa yang ditransformasi "
         f"menjadi {n_tfidf_features} fitur numerik menggunakan TF-IDF Vectorizer.", "Nominal\n(Kualitatif)"],
        ["3", "Label Perkembangan", "Y\n(Dependen)",
         "Predikat hasil evaluasi per aspek/mapel siswa: 'Baik', 'Cukup', atau 'Perlu Bimbingan'. "
         "Multi-output memprediksi Y untuk setiap aspek/mapel secara terpisah.", "Ordinal\n(Kualitatif)"],
    ]
    var_widths = [Inches(0.4), Inches(1.3), Inches(1.0), Inches(3.5), Inches(1.0)]
    build_table(doc, var_headers, var_data, var_widths, font_size=10,
                center_cols=[0, 2, 4])

    doc.add_paragraph()

    # =====================================================
    # SECTION 1.1: STRUKTUR VARIABEL SETELAH TRANSFORMASI
    # =====================================================
    h = doc.add_heading(level=2)
    r = h.add_run("1.1 Struktur Variabel Setelah Transformasi TF-IDF")
    fmt(r, size=12, bold=True, color=(44, 62, 80))

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(
        f"Setelah proses transformasi TF-IDF, setiap baris data memiliki total "
        f"{1 + n_tfidf_features} variabel input (fitur). Berikut struktur variabel per baris:"
    )
    fmt(r)

    # Structure table
    struct_headers = ["No", "Nama Fitur", "Sumber", "Tipe", "Keterangan"]
    struct_data = [
        ["1", "X1_Nilai", "X1 (Nilai)", "Numerik", "Nilai kuantitatif asli siswa (0-100)"],
        ["2", f"tfidf_1 ('{feature_names[0]}')", "X2 (Deskripsi)", "Numerik",
         f"Bobot TF-IDF kata '{feature_names[0]}'"],
        ["3", f"tfidf_2 ('{feature_names[1]}')", "X2 (Deskripsi)", "Numerik",
         f"Bobot TF-IDF kata '{feature_names[1]}'"],
        ["...", "...", "...", "...", "..."],
        [str(n_tfidf_features), f"tfidf_{n_tfidf_features} ('{feature_names[-2]}')",
         "X2 (Deskripsi)", "Numerik", f"Bobot TF-IDF kata '{feature_names[-2]}'"],
        [str(1 + n_tfidf_features), f"tfidf_{n_tfidf_features} ('{feature_names[-1]}')",
         "X2 (Deskripsi)", "Numerik", f"Bobot TF-IDF kata '{feature_names[-1]}'"],
    ]
    struct_widths = [Inches(0.4), Inches(2.0), Inches(1.3), Inches(0.8), Inches(2.5)]
    build_table(doc, struct_headers, struct_data, struct_widths, font_size=9,
                center_cols=[0, 3])

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(
        f"\nTotal variabel input per baris = 1 (X1_Nilai) + {n_tfidf_features} (TF-IDF) "
        f"= {1 + n_tfidf_features} fitur.\n"
        f"Berikut contoh matriks fitur gabungan untuk 10 data awal dan 10 data akhir dari data training "
        f"(N = {N_total} baris):"
    )
    fmt(r)

    # Feature matrix sample table (10 awal + 10 akhir)
    fm_headers = ["Data Ke-", "Nama Siswa", "Aspek/Mapel", "X1_Nilai",
                  f"tfidf_1\n({feature_names[0]})",
                  f"tfidf_2\n({feature_names[1]})",
                  f"tfidf_3\n({feature_names[2]})",
                  "...",
                  f"tfidf_{n_tfidf_features}\n({feature_names[-1]})",
                  "Non-Zero\nFeatures"]
    fm_data = []
    for i, row_idx in enumerate(subset_indices):
        s_row = subset.iloc[i]
        tfidf_row = X_tfidf_full[row_idx]
        nonzero_count = int(np.sum(tfidf_row > 0))

        fm_data.append([
            str(s_row['_orig_idx']),
            s_row['Nama Siswa'],
            str(s_row['Aspek / Mapel'])[:15],
            str(int(s_row['X1 (Nilai)'])),
            f"{tfidf_row[0]:.4f}",
            f"{tfidf_row[1]:.4f}",
            f"{tfidf_row[2]:.4f}",
            "...",
            f"{tfidf_row[-1]:.4f}",
            str(nonzero_count)
        ])

        # Add separator after row 10
        if i == 9:
            fm_data.append(["...", "...", "...", "...", "...", "...", "...", "...", "...", "..."])

    fm_widths = [Inches(0.5), Inches(1.0), Inches(1.0), Inches(0.6),
                 Inches(0.7), Inches(0.7), Inches(0.7), Inches(0.3), Inches(0.7), Inches(0.6)]
    build_table(doc, fm_headers, fm_data, fm_widths, font_size=8,
                center_cols=[0, 3, 4, 5, 6, 7, 8, 9])

    doc.add_paragraph()

    # =====================================================
    # SECTION 2: TF-IDF
    # =====================================================
    h = doc.add_heading(level=1)
    r = h.add_run("2. Proses Perhitungan TF-IDF Manual")
    fmt(r, size=14, bold=True, color=(44, 62, 80))

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(
        "Ekstraksi fitur TF-IDF (Term Frequency - Inverse Document Frequency) mengubah "
        "Deskripsi Capaian Pembelajaran (X2) menjadi representasi vektor numerik. "
        "Tahapan preprocessing meliputi: Case Folding, Cleansing (hapus tanda baca), "
        "Stopwords Removal (termasuk nama siswa dan istilah domain), dan Tokenisasi (kata > 2 karakter)."
    )
    fmt(r)

    # Equations
    p = doc.add_paragraph()
    r = p.add_run("Formula TF (Term Frequency):")
    fmt(r, italic=True)

    eq_tf = (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '  <m:r><m:t>TF(w, d) = </m:t></m:r>'
        '  <m:f>'
        '    <m:num><m:r><m:t>n(w, d)</m:t></m:r></m:num>'
        '    <m:den><m:r><m:t>Nd</m:t></m:r></m:den>'
        '  </m:f>'
        '</m:oMath>'
    )
    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_equation(p_eq, eq_tf)

    p = doc.add_paragraph()
    r = p.add_run(
        "Di mana n(w, d) = frekuensi kemunculan kata w pada dokumen d, "
        "Nd = total kata dalam dokumen d setelah preprocessing."
    )
    fmt(r, size=10, italic=True)

    p = doc.add_paragraph()
    r = p.add_run("Formula IDF (Inverse Document Frequency) menggunakan Logaritma Natural (ln):")
    fmt(r, italic=True)

    eq_idf = (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '  <m:r><m:t>IDF(w) = ln</m:t></m:r>'
        '  <m:d>'
        '    <m:dPr><m:begChr m:val="("/><m:endChr m:val=")"/></m:dPr>'
        '    <m:e>'
        '      <m:f>'
        '        <m:num><m:r><m:t>N</m:t></m:r></m:num>'
        '        <m:den><m:r><m:t>df(w)</m:t></m:r></m:den>'
        '      </m:f>'
        '    </m:e>'
        '  </m:d>'
        '</m:oMath>'
    )
    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_equation(p_eq, eq_idf)

    p = doc.add_paragraph()
    r = p.add_run(
        f"Di mana N = jumlah keseluruhan dokumen dalam korpus training (N = {N_total}) "
        f"dan df(w) = jumlah dokumen yang mengandung kata w."
    )
    fmt(r, size=10, italic=True)

    p = doc.add_paragraph()
    r = p.add_run("Formula pembobotan TF-IDF akhir:")
    fmt(r, italic=True)

    eq_tfidf = (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '  <m:r><m:t>TF-IDF(w, d) = TF(w, d) &#215; IDF(w)</m:t></m:r>'
        '</m:oMath>'
    )
    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_equation(p_eq, eq_tfidf)

    # Tracing example with first row
    first_row = subset.iloc[0]
    first_clean = first_row['clean_X2']
    first_words = first_clean.split()
    first_Nd = len(first_words)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(
        f"Sebagai contoh penelusuran manual pada Data Training ke-1 "
        f"(Siswa: {first_row['Nama Siswa']}, Aspek: {first_row['Aspek / Mapel']}):\n"
        f"1. Teks Asli: \"{str(first_row['X2 (Deskripsi Capaian)'])[:200]}...\"\n"
        f"2. Teks Preprocessed: \"{first_clean[:150]}...\"\n"
        f"3. Total Kata Bersih (Nd) = {first_Nd} kata."
    )
    fmt(r)

    # TF-IDF calculation table (top 3 words per 20 rows)
    p = doc.add_paragraph()
    r = p.add_run(
        "\nTabel berikut menyajikan perhitungan TF-IDF untuk 3 kata teratas "
        "dari 10 data awal dan 10 data akhir data training:"
    )
    fmt(r, italic=True)

    tfidf_headers = ["Data Ke-", "Nama Siswa", "Aspek", "Kata (w)", "Freq\n(nw,d)",
                     "Total\n(Nd)", "TF", "N", "df(w)", "IDF\n(ln)", "TF-IDF"]
    tfidf_rows = []
    prev_idx = None
    for item in tfidf_manual:
        # Add separator between first 10 and last 10
        if prev_idx is not None and prev_idx <= 10 and item['idx'] > 10:
            tfidf_rows.append(["...", "...", "...", "...", "...", "...",
                               "...", "...", "...", "...", "..."])
        tfidf_rows.append([
            str(item['idx']),
            item['nama'],
            str(item['aspek'])[:12],
            item['word'],
            str(item['freq']),
            str(item['Nd']),
            f"{item['tf']:.4f}",
            str(item['N']),
            str(item['df']),
            f"{item['idf']:.4f}",
            f"{item['tfidf']:.4f}"
        ])
        prev_idx = item['idx']

    tfidf_widths = [Inches(0.4), Inches(0.8), Inches(0.8), Inches(0.8), Inches(0.4),
                    Inches(0.4), Inches(0.5), Inches(0.3), Inches(0.4), Inches(0.5), Inches(0.5)]
    build_table(doc, tfidf_headers, tfidf_rows, tfidf_widths, font_size=8,
                center_cols=[0, 4, 5, 6, 7, 8, 9, 10])

    doc.add_paragraph()

    # =====================================================
    # SECTION 3: RANDOM FOREST
    # =====================================================
    h = doc.add_heading(level=1)
    r = h.add_run("3. Proses Klasifikasi Multi-Output Random Forest")
    fmt(r, size=14, bold=True, color=(44, 62, 80))

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(
        "Random Forest terdiri dari ensemble (kumpulan) pohon keputusan independen. "
        "Setiap pohon dibentuk menggunakan sampel bootstrap dan subset acak fitur. "
        "Prediksi akhir ditentukan oleh majority voting dari seluruh pohon."
    )
    fmt(r)

    # 3.1 Entropy & Gini
    h = doc.add_heading(level=2)
    r = h.add_run("3.1 Kriteria Pemecahan Simpul (Node Split)")
    fmt(r, size=12, bold=True, color=(127, 140, 141))

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run("Rumus Entropy untuk mengukur kemurnian data:")
    fmt(r)

    eq_ent = (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '  <m:r><m:t>Entropy(D) = -</m:t></m:r>'
        '  <m:nary>'
        '    <m:naryPr><m:chr m:val="&#8721;"/><m:limLoc m:val="undOvr"/></m:naryPr>'
        '    <m:sub><m:r><m:t>j=1</m:t></m:r></m:sub>'
        '    <m:sup><m:r><m:t>C</m:t></m:r></m:sup>'
        '    <m:e>'
        '      <m:sSub><m:e><m:r><m:t>p</m:t></m:r></m:e><m:sub><m:r><m:t>j</m:t></m:r></m:sub></m:sSub>'
        '      <m:r><m:t> log</m:t></m:r>'
        '      <m:sSub><m:e><m:r><m:t></m:t></m:r></m:e><m:sub><m:r><m:t>2</m:t></m:r></m:sub></m:sSub>'
        '      <m:d>'
        '        <m:dPr><m:begChr m:val="("/><m:endChr m:val=")"/></m:dPr>'
        '        <m:e><m:sSub><m:e><m:r><m:t>p</m:t></m:r></m:e><m:sub><m:r><m:t>j</m:t></m:r></m:sub></m:sSub></m:e>'
        '      </m:d>'
        '    </m:e>'
        '  </m:nary>'
        '</m:oMath>'
    )
    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_equation(p_eq, eq_ent)

    p = doc.add_paragraph()
    r = p.add_run("Rumus Gini Index:")
    fmt(r)

    eq_gini = (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '  <m:r><m:t>Gini(D) = 1 - </m:t></m:r>'
        '  <m:nary>'
        '    <m:naryPr><m:chr m:val="&#8721;"/><m:limLoc m:val="undOvr"/></m:naryPr>'
        '    <m:sub><m:r><m:t>j=1</m:t></m:r></m:sub>'
        '    <m:sup><m:r><m:t>C</m:t></m:r></m:sup>'
        '    <m:e>'
        '      <m:sSup>'
        '        <m:e><m:sSub><m:e><m:r><m:t>p</m:t></m:r></m:e><m:sub><m:r><m:t>j</m:t></m:r></m:sub></m:sSub></m:e>'
        '        <m:sup><m:r><m:t>2</m:t></m:r></m:sup>'
        '      </m:sSup>'
        '    </m:e>'
        '  </m:nary>'
        '</m:oMath>'
    )
    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_equation(p_eq, eq_gini)

    # 3.2 Voting table
    h = doc.add_heading(level=2)
    r = h.add_run("3.2 Perhitungan Multi-Output Voting")
    fmt(r, size=12, bold=True, color=(127, 140, 141))

    n_trees = len(rf.estimators_)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(
        f"Dalam Multi-Output Classifier, model memprediksi label perkembangan (Baik, Cukup, Perlu Bimbingan) "
        f"untuk setiap baris data. Model menggunakan {n_trees} pohon keputusan. "
        f"Setiap pohon memberikan suara independen, dan prediksi akhir ditentukan oleh mayoritas. "
        f"Berikut tabel voting untuk 10 data awal dan 10 data akhir:"
    )
    fmt(r)

    vote_headers = ["Data Ke-", "Nama Siswa", "Aspek/Mapel", "Y Aktual",
                    "Rincian Voting", "Prediksi Akhir"]
    vote_rows = []
    for i in range(len(subset)):
        s_row = subset.iloc[i]
        votes = preds_all_trees[:, i]
        unique_v, counts_v = np.unique(votes, return_counts=True)
        vote_dict = {le_label.classes_[int(c)]: int(v) for c, v in zip(unique_v, counts_v)}
        vote_str = ", ".join([f"{k} ({v})" for k, v in sorted(vote_dict.items(), key=lambda x: x[1], reverse=True)])
        final = le_label.classes_[int(np.argmax(np.bincount(votes.astype(int))))]

        vote_rows.append([
            str(s_row['_orig_idx']),
            s_row['Nama Siswa'],
            str(s_row['Aspek / Mapel'])[:18],
            s_row['Y (Label)'],
            vote_str,
            final
        ])

        if i == 9:
            vote_rows.append(["...", "...", "...", "...", "...", "..."])

    vote_widths = [Inches(0.5), Inches(1.0), Inches(1.2), Inches(1.0), Inches(2.5), Inches(1.0)]
    build_table(doc, vote_headers, vote_rows, vote_widths, font_size=8,
                center_cols=[0, 3, 5])

    doc.add_paragraph()

    # =====================================================
    # SECTION 4: HYPERPARAMETER TUNING
    # =====================================================
    h = doc.add_heading(level=1)
    r = h.add_run("4. Hyperparameter Tuning (HPT)")
    fmt(r, size=14, bold=True, color=(44, 62, 80))

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(
        "Hyperparameter Tuning menggunakan GridSearchCV dengan 3-fold cross-validation. "
        "Parameter yang dioptimasi: n_estimators, max_depth, dan min_samples_split. "
        "Berikut hasil iterasi pengujian (12 kombinasi total):"
    )
    fmt(r)

    # Show top 6 and best from real results
    hpt_headers = ["Rank", "n_estimators", "max_depth", "min_samples_split",
                   "Mean Accuracy", "Std", "Status"]
    hpt_rows = []
    top_results = cv_results.head(12)
    for _, cv_row in top_results.iterrows():
        rank = int(cv_row['rank_test_score'])
        n_est = int(cv_row['param_n_estimators'])
        m_depth = str(cv_row['param_max_depth'])
        m_split = int(cv_row['param_min_samples_split'])
        mean_acc = f"{cv_row['mean_test_score']:.4f}"
        std_acc = f"{cv_row['std_test_score']:.4f}"
        status = "TERBAIK (BEST)" if rank == 1 else ""
        hpt_rows.append([str(rank), str(n_est), m_depth, str(m_split),
                         mean_acc, std_acc, status])

    hpt_widths = [Inches(0.4), Inches(0.9), Inches(0.8), Inches(1.2),
                  Inches(1.0), Inches(0.8), Inches(1.2)]
    build_table(doc, hpt_headers, hpt_rows, hpt_widths, font_size=9,
                center_cols=[0, 1, 2, 3, 4, 5, 6],
                highlight_row_fn=lambda r, d: d[6] == "TERBAIK (BEST)")

    p = doc.add_paragraph()
    r = p.add_run(
        f"\nKesimpulan: Model akhir menggunakan konfigurasi terbaik: "
        f"n_estimators={best_params['n_estimators']}, "
        f"max_depth={best_params['max_depth']}, "
        f"min_samples_split={best_params['min_samples_split']} "
        f"dengan akurasi validasi silang = {best_score:.4f} ({best_score*100:.2f}%)."
    )
    fmt(r, italic=True)

    doc.add_paragraph()

    # =====================================================
    # SECTION 5: SHAP
    # =====================================================
    h = doc.add_heading(level=1)
    r = h.add_run("5. Penjelasan Interpretasi Model Menggunakan SHAP")
    fmt(r, size=14, bold=True, color=(44, 62, 80))

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(
        "SHAP (SHapley Additive exPlanations) sebagai Explainable AI (XAI) menjelaskan "
        "bagaimana model membuat keputusan secara individual. Nilai Shapley mewakili kontribusi "
        "masing-masing fitur terhadap probabilitas prediksi kelas tertentu."
    )
    fmt(r)

    p = doc.add_paragraph()
    r = p.add_run("Persamaan sifat Aditif SHAP:")
    fmt(r)

    eq_shap = (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        '  <m:r><m:t>f(x) = </m:t></m:r>'
        '  <m:sSub><m:e><m:r><m:t>&#966;</m:t></m:r></m:e>'
        '  <m:sub><m:r><m:t>0</m:t></m:r></m:sub></m:sSub>'
        '  <m:r><m:t> + </m:t></m:r>'
        '  <m:nary>'
        '    <m:naryPr><m:chr m:val="&#8721;"/><m:limLoc m:val="undOvr"/></m:naryPr>'
        '    <m:sub><m:r><m:t>i=1</m:t></m:r></m:sub>'
        '    <m:sup><m:r><m:t>M</m:t></m:r></m:sup>'
        '    <m:e><m:sSub><m:e><m:r><m:t>&#966;</m:t></m:r></m:e>'
        '    <m:sub><m:r><m:t>i</m:t></m:r></m:sub></m:sSub></m:e>'
        '  </m:nary>'
        '</m:oMath>'
    )
    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_equation(p_eq, eq_shap)

    p = doc.add_paragraph()
    r = p.add_run(
        f"Di mana f(x) = probabilitas prediksi, phi_0 = Base Value, "
        f"phi_i = kontribusi fitur ke-i. Total fitur M = {1 + n_tfidf_features}."
    )
    fmt(r, size=10, italic=True)

    all_feature_names = ['X1_Nilai'] + feature_names

    # SHAP table
    p = doc.add_paragraph()
    r = p.add_run(
        "\nTabel rincian SHAP untuk 10 data awal dan 10 data akhir:"
    )
    fmt(r, italic=True)

    shap_headers = ["Data Ke-", "Nama", "Prediksi", "Base Value\n(phi_0)",
                    "Sum phi_i", "f(x)", "Top 1 Fitur", "Kontribusi 1",
                    "Top 2 Fitur", "Kontribusi 2"]
    shap_rows = []

    for i in range(len(subset)):
        s_row = subset.iloc[i]
        pred_idx = int(rf.predict(X_subset[i:i + 1])[0])
        pred_label = le_label.classes_[pred_idx]

        # shap_values is list of arrays for multi-class
        if isinstance(shap_values, list):
            sv = shap_values[pred_idx][i]
            bv = base_values[pred_idx] if isinstance(base_values, (list, np.ndarray)) else base_values
        else:
            sv = shap_values[i, :, pred_idx] if shap_values.ndim == 3 else shap_values[i]
            bv = base_values[pred_idx] if isinstance(base_values, (list, np.ndarray)) else base_values

        sum_phi = float(np.sum(sv))
        fx = float(bv) + sum_phi

        # Top 2 features
        top_idx = np.argsort(np.abs(sv))[-2:][::-1]
        top1_name = all_feature_names[top_idx[0]] if top_idx[0] < len(all_feature_names) else f"f_{top_idx[0]}"
        top1_val = sv[top_idx[0]]
        top2_name = all_feature_names[top_idx[1]] if top_idx[1] < len(all_feature_names) else f"f_{top_idx[1]}"
        top2_val = sv[top_idx[1]]

        shap_rows.append([
            str(s_row['_orig_idx']),
            s_row['Nama Siswa'],
            pred_label,
            f"{bv:.4f}",
            f"{sum_phi:.4f}",
            f"{fx:.4f}",
            top1_name,
            f"{top1_val:+.4f}",
            top2_name,
            f"{top2_val:+.4f}"
        ])

        if i == 9:
            shap_rows.append(["...", "...", "...", "...", "...", "...",
                              "...", "...", "...", "..."])

    shap_widths = [Inches(0.4), Inches(0.8), Inches(0.7), Inches(0.7),
                   Inches(0.6), Inches(0.6), Inches(0.9), Inches(0.7),
                   Inches(0.9), Inches(0.7)]
    build_table(doc, shap_headers, shap_rows, shap_widths, font_size=8,
                center_cols=[0, 2, 3, 4, 5, 7, 9])

    # Save
    doc.save(OUTPUT_PATH)
    print(f"\nDocument saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    generate_docx()
